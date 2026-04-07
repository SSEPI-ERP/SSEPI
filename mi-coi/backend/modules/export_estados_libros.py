# backend/modules/export_estados_libros.py
"""Exportación PDF (ReportLab), Excel (openpyxl) y Word (python-docx) para estados y libros."""
from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Sequence, Tuple


def _logo_path() -> Optional[str]:
    p = os.path.abspath(
        os.path.join(os.path.dirname(__file__), "..", "..", "assets", "ssepi_logo.png")
    )
    return p if os.path.isfile(p) else None


def export_pdf_tabla(
    titulo: str,
    filas: Sequence[Sequence[Any]],
    encabezados: Sequence[str],
    path_pdf: str,
    *,
    empresa: str = "",
    subtitulo: str = "",
    rfc_empresa: str = "",
) -> Dict[str, Any]:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as e:
        return {"exito": False, "error": f"ReportLab: {e}"}
    try:
        doc = SimpleDocTemplate(path_pdf, pagesize=letter, topMargin=0.65 * inch, bottomMargin=0.55 * inch)
        story = []
        styles = getSampleStyleSheet()
        lp = _logo_path()
        if lp:
            try:
                img = Image(lp, width=1.1 * inch, height=0.35 * inch)
                story.append(img)
                story.append(Spacer(1, 6))
            except Exception:
                pass
        if empresa:
            story.append(Paragraph(f"<b>{empresa}</b>", styles["Title"]))
            story.append(Spacer(1, 4))
        story.append(Paragraph(f"<b>{titulo}</b>", styles["Heading2"]))
        if subtitulo:
            story.append(Paragraph(subtitulo, styles["Normal"]))
        story.append(Paragraph(f"<font size=9 color=grey>{datetime.now():%d/%m/%Y %H:%M}</font>", styles["Normal"]))
        story.append(Spacer(1, 12))
        data = [list(encabezados)] + [list(map(lambda x: f"{x:,.2f}" if isinstance(x, float) else str(x), r)) for r in filas]
        t = Table(data, repeatRows=1)
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a5f")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                ]
            )
        )
        story.append(t)
        doc.build(story)
        return {"exito": True, "ruta": path_pdf}
    except Exception as e:
        return {"exito": False, "error": str(e)}


def export_xlsx_tabla(
    titulo: str,
    filas: Sequence[Sequence[Any]],
    encabezados: Sequence[str],
    path_xlsx: str,
    *,
    empresa: str = "",
    hoja: str = "Reporte",
    rfc_empresa: str = "",
) -> Dict[str, Any]:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError as e:
        return {"exito": False, "error": f"openpyxl: {e}"}
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = hoja[:31]
        r = 1
        if empresa:
            ws.cell(r, 1, empresa).font = Font(bold=True, size=12)
            r += 1
        rfc = (rfc_empresa or "").strip()
        if rfc:
            ws.cell(r, 1, f"RFC: {rfc}").font = Font(size=10)
            r += 1
        ws.cell(r, 1, titulo).font = Font(bold=True, size=14)
        r += 2
        head_fill = PatternFill("solid", fgColor="1E3A5F")
        hf = Font(bold=True, color="FFFFFF")
        for c, h in enumerate(encabezados, start=1):
            cell = ws.cell(r, c, h)
            cell.fill = head_fill
            cell.font = hf
            cell.alignment = Alignment(horizontal="center")
        r += 1
        for fila in filas:
            for c, val in enumerate(fila, start=1):
                ws.cell(r, c, val)
            r += 1
        wb.save(path_xlsx)
        return {"exito": True, "ruta": path_xlsx}
    except Exception as e:
        return {"exito": False, "error": str(e)}


def export_docx_tabla(
    titulo: str,
    filas: Sequence[Sequence[Any]],
    encabezados: Sequence[str],
    path_docx: str,
    *,
    empresa: str = "",
    rfc_empresa: str = "",
) -> Dict[str, Any]:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as e:
        return {"exito": False, "error": f"python-docx: {e}"}
    try:
        doc = Document()
        lp = _logo_path()
        if lp:
            try:
                doc.add_picture(lp, width=Inches(1.4))
            except Exception:
                pass
        if empresa:
            p = doc.add_paragraph()
            r = p.add_run(empresa)
            r.bold = True
            r.font.size = Pt(14)
        rfc = (rfc_empresa or "").strip()
        if rfc:
            doc.add_paragraph(f"RFC: {rfc}")
        doc.add_heading(titulo, level=1)
        doc.add_paragraph(datetime.now().strftime("%d/%m/%Y %H:%M"))
        table = doc.add_table(rows=1 + len(filas), cols=len(encabezados))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(encabezados):
            hdr[i].text = str(h)
        for ri, fila in enumerate(filas, start=1):
            row = table.rows[ri].cells
            for j, val in enumerate(fila):
                row[j].text = f"{val:,.2f}" if isinstance(val, float) else str(val)
        doc.save(path_docx)
        return {"exito": True, "ruta": path_docx}
    except Exception as e:
        return {"exito": False, "error": str(e)}


def export_estado_resultados_pdf(datos: Dict[str, Any], path_pdf: str, *, empresa: str = "") -> Dict[str, Any]:
    if not datos.get("exito"):
        return {"exito": False, "error": datos.get("error", "Sin datos")}
    sec = datos.get("secciones") or {}
    filas = [(k.replace("_", " ").title(), float(v)) for k, v in sec.items()]
    return export_pdf_tabla(
        "Estado de resultados",
        filas,
        ("Concepto", "Importe"),
        path_pdf,
        empresa=empresa,
        subtitulo=f"{datos.get('fecha_inicio','')} — {datos.get('fecha_fin','')}",
    )


def export_dict_secciones_pdf(
    titulo: str,
    datos: Dict[str, Any],
    path_pdf: str,
    *,
    empresa: str = "",
    clave_secciones: str = "secciones",
) -> Dict[str, Any]:
    """Exporta un dict con subdict de concepto→importe (balance, flujo anidado plano, etc.)."""
    if not datos.get("exito"):
        return {"exito": False, "error": datos.get("error", "Sin datos")}
    sec = datos.get(clave_secciones)
    if sec is None and clave_secciones == "secciones" and "actividades_operacion" in datos:
        flat = []
        for k, v in (datos.get("actividades_operacion") or {}).items():
            flat.append((str(k), float(v or 0)))
        for k, v in (datos.get("actividades_financiamiento") or {}).items():
            flat.append((str(k), float(v or 0)))
        for k, v in (datos.get("reconciliacion_caja") or {}).items():
            flat.append((str(k), float(v or 0)))
        filas = flat
    elif isinstance(sec, dict):
        filas = [(k.replace("_", " ").title(), float(v or 0)) for k, v in sec.items()]
    else:
        return {"exito": False, "error": "Estructura no soportada"}
    return export_pdf_tabla(
        titulo,
        filas,
        ("Concepto", "Importe"),
        path_pdf,
        empresa=empresa,
        subtitulo=str(datos.get("fecha_corte") or datos.get("fecha_inicio") or "") + " — " + str(datos.get("fecha_fin") or ""),
    )


def export_libro_diario_pdf(diario: Dict[str, Any], path_pdf: str, *, empresa: str = "") -> Dict[str, Any]:
    if not diario.get("exito"):
        return {"exito": False, "error": diario.get("error", "Sin datos")}
    filas: List[Tuple[Any, ...]] = []
    for p in diario.get("polizas") or []:
        for ln in (p.get("lineas") or p.get("partidas") or []):
            cfdi = ln.get("cfdi") or {}
            uuid_s = (cfdi.get("uuid") or ln.get("uuid") or "")[:36]
            filas.append(
                (
                    p.get("fecha"),
                    p.get("tipo_poliza"),
                    p.get("numero_poliza"),
                    ln.get("num_cuenta"),
                    ln.get("cargo"),
                    ln.get("abono"),
                    uuid_s,
                )
            )
    return export_pdf_tabla(
        "Libro diario",
        filas,
        ("Fecha", "Tipo", "No.", "Cuenta", "Cargo", "Abono", "UUID"),
        path_pdf,
        empresa=empresa,
        subtitulo=f"{diario.get('fecha_inicio')} — {diario.get('fecha_fin')}",
    )
